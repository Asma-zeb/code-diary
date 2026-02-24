from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Create document
doc = Document()

# Title
title = doc.add_heading('Sohaib Khattak - GitHub Projects Portfolio', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Profile Section
doc.add_heading('Profile', level=1)
p = doc.add_paragraph()
p.add_run('GitHub: ').bold = True
p.add_run('github.com/Sohaib-Khattak\n')
p.add_run('Name: ').bold = True
p.add_run('Sohaib Khattak\n')
p.add_run('Location: ').bold = True
p.add_run('Peshawar, Charsadda\n')
p.add_run('Education: ').bold = True
p.add_run('BS Computer Science, Abdul Wali Khan University Mardan (AWKUM)')

# Skills Section
doc.add_heading('Skills & Certifications', level=1)
skills = [
    'Web Development',
    'Agentic & Robotics Engineering (Certified - PIAIC, CECOS University Peshawar)',
    'Python Programming',
    'Data Analytics with PowerBI',
    'ggplot2',
    'n8n Automation',
    'Machine Learning'
]
for skill in skills:
    doc.add_paragraph(skill, style='List Bullet')

# Repositories Overview
doc.add_heading('Repositories Overview', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Repository'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'Primary Technology'

for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

repos = [
    ('Learn-Log', 'First repository with ML projects and web dev', 'Jupyter Notebook'),
    ('khattak-automation-hub', 'n8n workflows and automation projects', 'n8n'),
    ('SKILLS', 'AI agent skill files repository', 'AI Agents'),
    ('Agentic-AI-Documentary-', 'Documentation hub for autonomous agents', 'Documentation')
]

for repo, desc, tech in repos:
    row_cells = table.add_row().cells
    row_cells[0].text = repo
    row_cells[1].text = desc
    row_cells[2].text = tech

# Projects Detail Section
doc.add_heading('Projects Detail', level=1)

# Learn-Log Repository
doc.add_heading('1. Learn-Log Repository', level=2)

doc.add_heading('Home Prices Predictions (Based on Living Area)', level=3)
p = doc.add_paragraph()
p.add_run('Type: ').bold = True
p.add_run('Machine Learning / Data Science\n')
p.add_run('Format: ').bold = True
p.add_run('Jupyter Notebook (.ipynb)\n')
p.add_run('Description: ').bold = True
p.add_run('A machine learning project that predicts home prices based on living area using regression techniques.')

doc.add_heading('Unsupervised Learning Methods for Alzheimer\'s Disease Data', level=3)
p = doc.add_paragraph()
p.add_run('Type: ').bold = True
p.add_run('Research / Machine Learning\n')
p.add_run('Format: ').bold = True
p.add_run('Jupyter Notebook (.ipynb) + PDF Document\n')
p.add_run('Description: ').bold = True
p.add_run('Research project applying unsupervised learning methods for functional data with missing observations, specifically focused on Alzheimer\'s disease data analysis.')

doc.add_heading('Quiz Game', level=3)
p = doc.add_paragraph()
p.add_run('Type: ').bold = True
p.add_run('Web Development\n')
p.add_run('Format: ').bold = True
p.add_run('HTML, CSS, JavaScript\n')
p.add_run('Description: ').bold = True
p.add_run('An interactive quiz game web application with styled UI and game logic.')

# khattak-automation-hub Repository
doc.add_heading('2. khattak-automation-hub Repository', level=2)

doc.add_heading('Automated AI Assistant for Chat, Data & Email', level=3)
p = doc.add_paragraph()
p.add_run('Type: ').bold = True
p.add_run('Automation Workflow\n')
p.add_run('Format: ').bold = True
p.add_run('n8n Workflow (.json)\n')
p.add_run('Description: ').bold = True
p.add_run('An AI-powered automation workflow that handles chat interactions, data processing, and email management automatically.')

doc.add_heading('Automated NASA Solar Flare Data Integration Workflow', level=3)
p = doc.add_paragraph()
p.add_run('Type: ').bold = True
p.add_run('Automation / Data Integration\n')
p.add_run('Format: ').bold = True
p.add_run('n8n Workflow (.json)\n')
p.add_run('Description: ').bold = True
p.add_run('Integrates NASA solar flare data into automated workflows for real-time space weather data processing.')

doc.add_heading('Weekly Sales Report Workflow', level=3)
p = doc.add_paragraph()
p.add_run('Type: ').bold = True
p.add_run('Business Automation\n')
p.add_run('Format: ').bold = True
p.add_run('n8n Workflow (.json)\n')
p.add_run('Description: ').bold = True
p.add_run('Automated workflow for generating and distributing weekly sales reports.')

# SKILLS Repository
doc.add_heading('3. SKILLS Repository', level=2)

doc.add_heading('Claude Code Skills Lab', level=3)
p = doc.add_paragraph()
p.add_run('Type: ').bold = True
p.add_run('AI Agent Skills Development\n')
p.add_run('Format: ').bold = True
p.add_run('Skill Files / Configuration\n')
p.add_run('Location: ').bold = True
p.add_run('claude-code/claude-code-skills-lab-main/\n')
p.add_run('Description: ').bold = True
p.add_run('A centralized collection of skill files created for AI agents across multiple domains and use cases, specifically for Claude Code.')

# Agentic-AI-Documentary Repository
doc.add_heading('4. Agentic-AI-Documentary- Repository', level=2)

doc.add_heading('Agentic AI Documentary (Knowledge Hub)', level=3)
p = doc.add_paragraph()
p.add_run('Type: ').bold = True
p.add_run('Documentation / Knowledge Base\n')
p.add_run('Format: ').bold = True
p.add_run('Markdown Documentation\n')
p.add_run('Description: ').bold = True
p.add_run('A structured knowledge hub dedicated to documenting:\n')
doc.add_paragraph('Research on CLI-based autonomous agent systems', style='List Bullet')
doc.add_paragraph('Architectures and implementations', style='List Bullet')
doc.add_paragraph('Case studies', style='List Bullet')
doc.add_paragraph('Skills engineering', style='List Bullet')
doc.add_paragraph('MCP (Model Context Protocol) development', style='List Bullet')
doc.add_paragraph('Multi-agent automation frameworks', style='List Bullet')

# Project Statistics
doc.add_heading('Project Statistics', level=1)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Category'
hdr_cells[1].text = 'Count'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

stats = [
    ('Total Repositories', '4'),
    ('Total Projects', '8'),
    ('Machine Learning Projects', '2'),
    ('Automation Workflows (n8n)', '3'),
    ('Web Development Projects', '1'),
    ('AI Agent Skills', '1'),
    ('Documentation/Knowledge Base', '1')
]

for category, count in stats:
    row_cells = table.add_row().cells
    row_cells[0].text = category
    row_cells[1].text = count

# Technology Stack
doc.add_heading('Technology Stack', level=1)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Technology'
hdr_cells[1].text = 'Usage'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

tech_stack = [
    ('Python', 'ML, Data Analysis'),
    ('Jupyter Notebook', 'ML Projects, Research'),
    ('n8n', 'Automation Workflows'),
    ('HTML/CSS/JavaScript', 'Web Development'),
    ('PowerBI', 'Data Analytics'),
    ('AI Agents', 'Skills Development')
]

for tech, usage in tech_stack:
    row_cells = table.add_row().cells
    row_cells[0].text = tech
    row_cells[1].text = usage

# GitHub Activity
doc.add_heading('GitHub Activity', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Repository'
hdr_cells[1].text = 'Commits'
hdr_cells[2].text = 'Stars'
hdr_cells[3].text = 'Forks'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

activity = [
    ('Learn-Log', '-', '-', '-'),
    ('khattak-automation-hub', '12', '1', '0'),
    ('SKILLS', '6', '1', '0'),
    ('Agentic-AI-Documentary-', '2', '-', '-')
]

for repo, commits, stars, forks in activity:
    row_cells = table.add_row().cells
    row_cells[0].text = repo
    row_cells[1].text = commits
    row_cells[2].text = stars
    row_cells[3].text = forks

# Contact Section
doc.add_heading('Contact & Links', level=1)
p = doc.add_paragraph()
p.add_run('GitHub: ').bold = True
p.add_run('https://github.com/Sohaib-Khattak\n')
p.add_run('Location: ').bold = True
p.add_run('Peshawar, Charsadda, Pakistan')

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Document Generated: February 23, 2026').italic = True

# Save document
doc.save('Sohaib_GitHub_Projects.docx')
print('Document saved successfully as Sohaib_GitHub_Projects.docx')
